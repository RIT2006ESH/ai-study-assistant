from docx import Document
from typing import Optional


async def process_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If DOCX processing fails
    """
    try:
        doc = Document(file_path)
        
        # Extract text from all paragraphs
        extracted_text = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                extracted_text.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        extracted_text.append(text)
        
        # Join all text with newlines
        full_text = "\n\n".join(extracted_text)
        
        if not full_text.strip():
            raise ValueError("No text could be extracted from the DOCX file")
        
        return full_text
        
    except Exception as e:
        raise Exception(f"DOCX processing failed: {str(e)}")


async def get_docx_metadata(file_path: str) -> dict:
    """
    Extract metadata from a DOCX file
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Dictionary containing DOCX metadata
    """
    try:
        doc = Document(file_path)
        
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
        
        # Try to get core properties
        try:
            core_props = doc.core_properties
            metadata["title"] = core_props.title
            metadata["author"] = core_props.author
            metadata["subject"] = core_props.subject
        except:
            pass
        
        return metadata
        
    except Exception as e:
        return {"error": str(e)}
